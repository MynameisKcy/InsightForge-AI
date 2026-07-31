"""
Export Agent: Markdown → Word/PDF/HTML/Markdown 导出。
"""

import os
import re
import sys
from datetime import datetime
from typing import Any

CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(CURRENT_DIR)
for path in (PROJECT_ROOT, os.path.dirname(PROJECT_ROOT)):
    if path not in sys.path:
        sys.path.insert(0, path)

from agents.base import BaseAgent
from utils.logger_handler import logger
from utils.path_tool import get_abs_path

# 延迟导入，优雅降级
_docx_available = True
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    _docx_available = False
    logger.warning("python-docx not installed. Word export disabled.")

_pdf_available = True
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    _pdf_available = False
    logger.warning("reportlab not installed. PDF export disabled.")

_cjk_font_registered = False
_cjk_font_name = None  # 注册成功后的字体名，None 表示未注册

EXPORT_DIR = "reports"


def _ensure_export_dir(subdir: str = "") -> str:
    path = get_abs_path(os.path.join(EXPORT_DIR, subdir))
    os.makedirs(path, exist_ok=True)
    return path


class ExportAgent(BaseAgent):
    """导出 Agent：将 Markdown 报告导出为 Word、PDF、HTML、Markdown 文件。"""

    name = "export_agent"

    def run(self, input_data: dict) -> dict:
        """
        input_data = {
            "markdown": str,
            "title": str,
            "formats": ["md", "docx", "pdf", "html"],  # 默认所有格式
        }
        returns: {"files": [{"format": "docx", "path": str}], "errors": []}
        """
        markdown = input_data.get("markdown", "")
        title = input_data.get("title", "数据分析报告")
        formats = input_data.get("formats", ["md", "docx", "pdf", "html"])

        if not markdown:
            return {"files": [], "errors": ["No markdown content to export"]}

        results = {"files": [], "errors": []}

        for fmt in formats:
            try:
                path = self._export(markdown, title, fmt)
                if path:
                    results["files"].append({"format": fmt, "path": path})
                else:
                    results["errors"].append(f"Failed to export {fmt}")
            except Exception as e:
                logger.error(f"Export to {fmt} failed: {e}")
                results["errors"].append(f"Export {fmt} failed: {e}")

        return results

    def _register_cjk_font(self) -> str | None:
        """注册一个 Windows 自带的中文字体供 reportlab 使用。

        优先 msyh.ttc（微软雅黑）→ simsun.ttc（宋体），用 subfontIndex=0 取集合首字体。
        全局只注册一次。都找不到则记 warning 返回 None，不阻断 PDF 生成。
        """
        global _cjk_font_registered, _cjk_font_name
        if _cjk_font_registered:
            return _cjk_font_name
        if not _pdf_available:
            _cjk_font_registered = True
            return None

        fonts_dir = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts")
        candidates = ["msyh.ttc", "msyhbd.ttc", "simsun.ttc", "simhei.ttf"]
        for fname in candidates:
            fpath = os.path.join(fonts_dir, fname)
            if os.path.exists(fpath):
                try:
                    pdfmetrics.registerFont(TTFont("CJK", fpath, subfontIndex=0))
                    _cjk_font_name = "CJK"
                    logger.info(f"Registered CJK font: {fpath}")
                    break
                except Exception as e:
                    logger.warning(f"Failed to register font {fpath}: {e}")
        if _cjk_font_name is None:
            logger.warning("No CJK font found; PDF Chinese may render incorrectly.")
        _cjk_font_registered = True
        return _cjk_font_name

    def _export(self, markdown: str, title: str, fmt: str) -> str | None:
        """按格式导出。"""
        fmt = fmt.lower().strip()
        if fmt in ("md", "markdown"):
            return self._export_markdown(markdown, title)
        elif fmt in ("docx", "word"):
            return self._export_docx(markdown, title)
        elif fmt == "pdf":
            return self._export_pdf(markdown, title)
        elif fmt == "html":
            return self._export_html(markdown, title)
        else:
            logger.warning(f"Unknown export format: {fmt}")
            return None

    def _export_markdown(self, markdown: str, title: str) -> str:
        """导出为 .md 文件（直接写入）。"""
        output_dir = _ensure_export_dir("markdown")
        filename = _make_filename(title, "md")
        filepath = os.path.join(output_dir, filename)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(markdown)
        logger.info(f"Markdown exported to {filepath}")
        return filepath

    def _export_docx(self, markdown: str, title: str) -> str | None:
        """导出为 Word (.docx) 文件。"""
        if not _docx_available:
            logger.warning("python-docx not available, skipping Word export")
            return None

        doc = Document()

        # 标题
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 按行解析 Markdown
        lines = markdown.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            if not line:
                doc.add_paragraph("")
                i += 1
                continue

            # 标题
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            # 表格
            elif line.startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                i -= 1  # 回退，因为循环会 +1
                self._add_docx_table(doc, table_lines)
            # 图片引用
            elif line.startswith("!["):
                match = re.match(r"!\[.*?\]\((.*?)\)", line)
                if match:
                    img_path = match.group(1)
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(5.5))
            # 普通段落
            else:
                # 移除 Markdown 标记
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                clean = re.sub(r"`(.+?)`", r"\1", clean)
                if clean.startswith("- "):
                    doc.add_paragraph(clean[2:], style="List Bullet")
                elif re.match(r"^\d+\.\s", clean):
                    doc.add_paragraph(re.sub(r"^\d+\.\s", "", clean), style="List Number")
                else:
                    doc.add_paragraph(clean)
            i += 1

        output_dir = _ensure_export_dir("word")
        filename = _make_filename(title, "docx")
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        logger.info(f"Word exported to {filepath}")
        return filepath

    def _add_docx_table(self, doc, table_lines: list[str]):
        """在 Word 文档中添加表格。"""
        if len(table_lines) < 2:
            return
        # 解析表头
        header_cells = [c.strip() for c in table_lines[0].split("|") if c.strip()]
        # 跳过分隔行
        data_start = 2 if len(table_lines) > 1 and "---" in table_lines[1] else 1
        rows = []
        for line in table_lines[data_start:]:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if cells:
                rows.append(cells)

        if header_cells:
            table = doc.add_table(rows=1 + len(rows), cols=len(header_cells), style="Light List Accent 1")
            # 表头
            for j, cell_text in enumerate(header_cells):
                table.rows[0].cells[j].text = cell_text
            # 数据行
            for i, row in enumerate(rows):
                for j in range(min(len(row), len(header_cells))):
                    table.rows[i + 1].cells[j].text = row[j]
            doc.add_paragraph("")

    def _export_pdf(self, markdown: str, title: str) -> str | None:
        """导出为 PDF 文件。"""
        if not _pdf_available:
            logger.warning("reportlab not available, skipping PDF export")
            return None

        output_dir = _ensure_export_dir("pdf")
        filename = _make_filename(title, "pdf")
        filepath = os.path.join(output_dir, filename)

        cjk = self._register_cjk_font()
        font_for_pdf = cjk or "Helvetica"

        doc = SimpleDocTemplate(filepath, pagesize=A4,
                                leftMargin=20*mm, rightMargin=20*mm,
                                topMargin=20*mm, bottomMargin=20*mm)
        story = []

        # 使用基本样式
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("CustomTitle", parent=styles["Title"],
                                     fontName=font_for_pdf, fontSize=18, spaceAfter=12)
        h1_style = ParagraphStyle("CustomH1", parent=styles["Heading1"],
                                  fontName=font_for_pdf, fontSize=16, spaceBefore=12, spaceAfter=6)
        h2_style = ParagraphStyle("CustomH2", parent=styles["Heading2"],
                                  fontName=font_for_pdf, fontSize=13, spaceBefore=10, spaceAfter=4)
        body_style = ParagraphStyle("CustomBody", parent=styles["Normal"],
                                    fontName=font_for_pdf, fontSize=10, leading=14, spaceAfter=6)

        styles_dict = {"font_name": font_for_pdf}

        # 添加标题
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 6*mm))

        # 按行解析
        lines = markdown.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            if not line:
                story.append(Spacer(1, 3*mm))
                i += 1
                continue

            clean = self._clean_md_for_pdf(line)
            if not clean:
                i += 1
                continue

            if line.startswith("# "):
                story.append(Paragraph(clean, h1_style))
            elif line.startswith("## "):
                story.append(Paragraph(clean, h2_style))
            elif line.startswith("### "):
                story.append(Paragraph(clean, styles["Heading3"]))
            elif line.startswith("- "):
                story.append(Paragraph(f"• {clean}", body_style))
            elif line.startswith("---"):
                story.append(Spacer(1, 4*mm))
            elif line.startswith("!["):
                story.append(Paragraph("（图表见 HTML 版报告）", body_style))
            elif line.startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                i -= 1
                self._add_pdf_table(story, table_lines, styles_dict)
            else:
                story.append(Paragraph(clean, body_style))
            i += 1

        try:
            doc.build(story)
            logger.info(f"PDF exported to {filepath}")
            return filepath
        except Exception as e:
            logger.error(f"PDF build failed: {e}")
            return None

    def _add_pdf_table(self, story, table_lines: list, styles: dict):
        """把 markdown 表格行渲染为 reportlab Table。"""
        if len(table_lines) < 2:
            return
        header_cells = [c.strip() for c in table_lines[0].split("|") if c.strip()]
        data_start = 2 if len(table_lines) > 1 and "---" in table_lines[1] else 1
        rows = []
        for line in table_lines[data_start:]:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if cells:
                rows.append(cells)
        if not header_cells:
            return
        # 列数对齐
        col_count = len(header_cells)
        table_data = [header_cells] + [r[:col_count] + [""] * (col_count - len(r)) for r in rows]
        font_name = styles.get("font_name", "Helvetica")
        style = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2f2f2")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fafafa")]),
        ])
        tbl = Table(table_data, style=style, hAlign="LEFT")
        story.append(tbl)
        story.append(Spacer(1, 4 * mm))

    def _clean_md_for_pdf(self, text: str) -> str:
        """清理 Markdown 标记，转为纯文本 + HTML 标签（reportlab 支持的）。"""
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
        text = re.sub(r"`(.+?)`", r"<font face='Courier'>\1</font>", text)
        # 移除标题标记
        text = re.sub(r"^#{1,3}\s+", "", text)
        return text

    def _export_html(self, markdown: str, title: str) -> str:
        """导出为 HTML 文件（用基本 Markdown→HTML 转换）。"""
        output_dir = _ensure_export_dir("html")
        filename = _make_filename(title, "html")
        filepath = os.path.join(output_dir, filename)

        html = self._md_to_html(markdown, title)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(html)
        logger.info(f"HTML exported to {filepath}")
        return filepath

    def _md_to_html(self, markdown: str, title: str) -> str:
        """将 Markdown 转为基本 HTML。"""
        lines = markdown.split("\n")
        html_lines = [
            "<!DOCTYPE html>",
            "<html><head><meta charset='utf-8'>",
            f"<title>{title}</title>",
            "<style>",
            "body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }",
            "h1 { border-bottom: 2px solid #333; padding-bottom: 8px; }",
            "h2 { border-bottom: 1px solid #ddd; padding-bottom: 4px; }",
            "table { border-collapse: collapse; width: 100%; margin: 10px 0; }",
            "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "th { background-color: #f2f2f2; }",
            "img { max-width: 100%; }",
            "</style></head><body>",
        ]

        in_table = False
        in_thead = False
        skip_next = False

        for i, line in enumerate(lines):
            stripped = line.rstrip()

            if skip_next:
                skip_next = False
                continue

            if not stripped:
                if in_table:
                    html_lines.append("</tbody></table>")
                    in_table = False
                html_lines.append("")
                continue

            # 标题
            if stripped.startswith("# "):
                if in_table:
                    html_lines.append("</tbody></table>")
                    in_table = False
                html_lines.append(f"<h1>{stripped[2:]}</h1>")
            elif stripped.startswith("## "):
                if in_table:
                    html_lines.append("</tbody></table>")
                    in_table = False
                html_lines.append(f"<h2>{stripped[3:]}</h2>")
            elif stripped.startswith("### "):
                if in_table:
                    html_lines.append("</tbody></table>")
                    in_table = False
                html_lines.append(f"<h3>{stripped[4:]}</h3>")
            # 表格
            elif stripped.startswith("|"):
                if not in_table:
                    html_lines.append("<table><thead>")
                    in_table = True
                    in_thead = True
                if in_thead:
                    cells = [c.strip() for c in stripped.split("|") if c.strip()]
                    if all("---" in c for c in cells):
                        html_lines.append("</thead><tbody>")
                        in_thead = False
                        continue
                    cleaned = [self._html_clean(c) for c in cells]
                    html_lines.append(f"<tr><th>{'</th><th>'.join(cleaned)}</th></tr>")
                else:
                    cells = [c.strip() for c in stripped.split("|") if c.strip()]
                    cleaned = [self._html_clean(c) for c in cells]
                    html_lines.append(f"<tr><td>{'</td><td>'.join(cleaned)}</td></tr>")
            # 图片
            elif stripped.startswith("!["):
                match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
                if match:
                    alt, src = match.group(1), match.group(2)
                    html_lines.append(f'<img src="{src}" alt="{alt}">')
            # 分隔线
            elif stripped == "---":
                html_lines.append("<hr>")
            # 普通段落 / 列表
            else:
                if in_table:
                    html_lines.append("</tbody></table>")
                    in_table = False
                clean = self._html_clean(stripped)
                clean = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", clean)
                clean = re.sub(r"\*(.+?)\*", r"<em>\1</em>", clean)
                clean = re.sub(r"`(.+?)`", r"<code>\1</code>", clean)
                if stripped.startswith("- "):
                    html_lines.append(f"<li>{clean[2:]}</li>")
                elif re.match(r"^\d+\.\s", stripped):
                    numbered_clean = re.sub(r'^\d+\.\s', '', clean)
                    html_lines.append(f"<li>{numbered_clean}</li>")
                else:
                    html_lines.append(f"<p>{clean}</p>")

        if in_table:
            html_lines.append("</tbody></table>")

        html_lines.append("</body></html>")
        return "\n".join(html_lines)

    @staticmethod
    def _html_clean(text: str) -> str:
        """HTML 转义。"""
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _make_filename(title: str, ext: str) -> str:
    """生成安全文件名。"""
    safe = re.sub(r"[^\w\s-]", "", title)
    safe = re.sub(r"[-\s]+", "_", safe.strip())[:50]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{safe}_{timestamp}.{ext}"
