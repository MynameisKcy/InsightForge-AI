"""
Export Agent: Markdown → Word/PDF/HTML/Markdown 导出。

Markdown 方言的解析统一走 agents/markdown_blocks.py（块解析器，纯函数）；
本文件只保留各格式的渲染 adapter：Block → python-docx / reportlab / HTML。
行内 **粗**/*斜*/`码` 的转换语义各格式不同，由本文件的 _strip_inline_md /
_pdf_inline / _md_to_html 各自负责——解析器不做行内转换。
"""

import base64
import os
import re
from datetime import datetime

from agents.base import BaseAgent
from agents.markdown_blocks import (
    Blank,
    Heading,
    HorizontalRule,
    Image as MdImage,
    ListItem,
    Paragraph as MdParagraph,
    Table as MdTable,
    parse_markdown_blocks,
)
from utils.logger_handler import logger
from utils.path_tool import get_abs_path

# 延迟导入，优雅降级
_docx_available = True
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor  # noqa: F401  # 可用性探测导入
except ImportError:
    _docx_available = False
    logger.warning("python-docx not installed. Word export disabled.")

_pdf_available = True
try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT  # noqa: F401  # 可用性探测导入
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import (
        PageBreak,  # noqa: F401  # 可用性探测导入
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
except ImportError:
    _pdf_available = False
    logger.warning("reportlab not installed. PDF export disabled.")

_cjk_font_registered = False
_cjk_font_name = None  # 注册成功后的字体名，None 表示未注册

EXPORT_DIR = "reports"


def _ensure_export_dir(subdir: str = "") -> str:
    path = get_abs_path(os.path.join(EXPORT_DIR, subdir))
    os.makedirs(path, exist_ok=True)
    return path


class ExportAgent(BaseAgent):
    """导出 Agent：将 Markdown 报告导出为 Word、PDF、HTML、Markdown 文件。"""

    name = "export_agent"

    def run(self, input_data: dict) -> dict:
        """
        input_data = {
            "markdown": str,
            "title": str,
            "formats": ["md", "docx", "pdf", "html"],  # 默认所有格式
        }
        returns: {"files": [{"format": "docx", "path": str}], "errors": []}
        """
        markdown = input_data.get("markdown", "")
        title = input_data.get("title", "数据分析报告")
        formats = input_data.get("formats", ["md", "docx", "pdf", "html"])

        if not markdown:
            return {"files": [], "errors": ["No markdown content to export"]}

        results = {"files": [], "errors": []}

        for fmt in formats:
            try:
                path = self._export(markdown, title, fmt)
                if path:
                    results["files"].append({"format": fmt, "path": path})
                else:
                    results["errors"].append(f"Failed to export {fmt}")
            except Exception as e:
                logger.error(f"Export to {fmt} failed: {e}")
                results["errors"].append(f"Export {fmt} failed: {e}")

        return results

    def _register_cjk_font(self) -> str | None:
        """注册一个 Windows 自带的中文字体供 reportlab 使用。

        优先 msyh.ttc（微软雅黑）→ simsun.ttc（宋体），用 subfontIndex=0 取集合首字体。
        全局只注册一次。都找不到则记 warning 返回 None，不阻断 PDF 生成。
        """
        global _cjk_font_registered, _cjk_font_name
        if _cjk_font_registered:
            return _cjk_font_name
        if not _pdf_available:
            _cjk_font_registered = True
            return None

        fonts_dir = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts")
        candidates = ["msyh.ttc", "msyhbd.ttc", "simsun.ttc", "simhei.ttf"]
        for fname in candidates:
            fpath = os.path.join(fonts_dir, fname)
            if os.path.exists(fpath):
                try:
                    pdfmetrics.registerFont(TTFont("CJK", fpath, subfontIndex=0))
                    _cjk_font_name = "CJK"
                    logger.info(f"Registered CJK font: {fpath}")
                    break
                except Exception as e:
                    logger.warning(f"Failed to register font {fpath}: {e}")
        if _cjk_font_name is None:
            logger.warning("No CJK font found; PDF Chinese may render incorrectly.")
        _cjk_font_registered = True
        return _cjk_font_name

    def _export(self, markdown: str, title: str, fmt: str) -> str | None:
        """按格式导出。"""
        fmt = fmt.lower().strip()
        if fmt in ("md", "markdown"):
            return self._export_markdown(markdown, title)
        elif fmt in ("docx", "word"):
            return self._export_docx(markdown, title)
        elif fmt == "pdf":
            return self._export_pdf(markdown, title)
        elif fmt == "html":
            return self._export_html(markdown, title)
        else:
            logger.warning(f"Unknown export format: {fmt}")
            return None

    # ── 图表图片解析（Plotly 图表为交互式 HTML，导出需栅格 PNG）──

    def _resolve_chart_image(self, ref: str) -> str | None:
        """把 markdown 图片引用解析为本地 PNG 文件路径；无法解析返回 None。

        ref 可能是 Web URL（/reports/charts/foo.png|.html）、FS 绝对路径、或占位符文本。
        .html 引用会查找同名 .png（kaleido 在图表生成时产出的栅格图）。
        """
        if not ref:
            return None
        local = ref
        if ref.startswith("/reports/"):
            local = get_abs_path(ref.lstrip("/"))
        if not local or not os.path.exists(local):
            return None
        low = local.lower()
        if low.endswith(".png"):
            return local
        if low.endswith(".html"):
            png = os.path.splitext(local)[0] + ".png"
            return png if os.path.exists(png) else None
        return None

    def _png_data_uri(self, png_path: str) -> str:
        """PNG 文件 -> base64 data URI（供 HTML/MD 自包含嵌入）。"""
        with open(png_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("ascii")
        return f"data:image/png;base64,{b64}"

    def _scaled_image(self, png_path: str, max_w_mm: float = 160, max_h_mm: float = 110):
        """reportlab Image 流对象，按原始宽高保比缩放到 max_w×max_h 框内。"""
        from PIL import Image as PILImage
        with PILImage.open(png_path) as im:
            iw, ih = im.size
        max_w = max_w_mm * mm
        max_h = max_h_mm * mm
        scale = min(max_w / iw, max_h / ih)
        return RLImage(png_path, width=iw * scale, height=ih * scale)

    def _export_markdown(self, markdown: str, title: str) -> str:
        """导出为 .md 文件（图表图片内联为 base64 data URI，自包含可离线查看）。"""
        output_dir = _ensure_export_dir("markdown")
        filename = _make_filename(title, "md")
        filepath = os.path.join(output_dir, filename)
        content = self._inline_images_as_data_uri(markdown)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info(f"Markdown exported to {filepath}")
        return filepath

    def _inline_images_as_data_uri(self, markdown: str) -> str:
        """把 markdown 图表图片引用替换为 base64 data URI；无法解析的保留原样。"""
        def repl(m):
            alt, ref = m.group(1), m.group(2)
            png = self._resolve_chart_image(ref)
            if png:
                try:
                    return f"![{alt}]({self._png_data_uri(png)})"
                except Exception as e:
                    logger.warning(f"MD image inline failed ({png}): {e}")
            return m.group(0)
        return re.sub(r"!\[(.*?)\]\((.*?)\)", repl, markdown)

    def _export_docx(self, markdown: str, title: str) -> str | None:
        """导出为 Word (.docx) 文件。"""
        if not _docx_available:
            logger.warning("python-docx not available, skipping Word export")
            return None

        doc = Document()

        # 标题
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for block in parse_markdown_blocks(markdown):
            self._render_docx_block(doc, block)

        output_dir = _ensure_export_dir("word")
        filename = _make_filename(title, "docx")
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        logger.info(f"Word exported to {filepath}")
        return filepath

    @staticmethod
    def _strip_inline_md(text: str) -> str:
        """剥离行内 Markdown 标记（Word 走纯文本，不加粗/斜体 run）。"""
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        return text

    def _render_docx_block(self, doc, block) -> None:
        """单个 Block → Word 元素。"""
        if isinstance(block, Blank):
            doc.add_paragraph("")
        elif isinstance(block, Heading):
            doc.add_heading(block.text, level=block.level)
        elif isinstance(block, MdTable):
            self._add_docx_table(doc, block)
        elif isinstance(block, HorizontalRule):
            self._add_docx_hr(doc)
        elif isinstance(block, MdImage):
            png = self._resolve_chart_image(block.ref)
            if png:
                try:
                    doc.add_picture(png, width=Inches(5.5))
                except Exception as e:
                    logger.warning(f"docx add_picture failed ({png}): {e}")
                    doc.add_paragraph(f"（图表：{block.alt or '图'}，无法嵌入）")
            else:
                doc.add_paragraph(f"（图表：{block.alt or '图'}）")
        elif isinstance(block, ListItem):
            style = "List Number" if block.ordered else "List Bullet"
            doc.add_paragraph(self._strip_inline_md(block.text), style=style)
        else:  # MdParagraph
            doc.add_paragraph(self._strip_inline_md(block.text))

    @staticmethod
    def _add_docx_hr(doc):
        """在 Word 中画一条水平分隔线（段落底边框实现）。"""
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def _add_docx_table(self, doc, table: MdTable):
        """在 Word 文档中添加表格。"""
        header_cells = table.header
        if not header_cells:
            return
        table_obj = doc.add_table(rows=1 + len(table.rows), cols=len(header_cells), style="Light List Accent 1")
        # 表头
        for j, cell_text in enumerate(header_cells):
            table_obj.rows[0].cells[j].text = cell_text
        # 数据行
        for i, row in enumerate(table.rows):
            for j in range(min(len(row), len(header_cells))):
                table_obj.rows[i + 1].cells[j].text = row[j]
        doc.add_paragraph("")

    def _export_pdf(self, markdown: str, title: str) -> str | None:
        """导出为 PDF 文件。"""
        if not _pdf_available:
            logger.warning("reportlab not available, skipping PDF export")
            return None

        output_dir = _ensure_export_dir("pdf")
        filename = _make_filename(title, "pdf")
        filepath = os.path.join(output_dir, filename)

        cjk = self._register_cjk_font()
        font_for_pdf = cjk or "Helvetica"

        doc = SimpleDocTemplate(filepath, pagesize=A4,
                                leftMargin=20*mm, rightMargin=20*mm,
                                topMargin=20*mm, bottomMargin=20*mm)
        story = []

        # 使用基本样式
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("CustomTitle", parent=styles["Title"],
                                     fontName=font_for_pdf, fontSize=18, spaceAfter=12)
        h1_style = ParagraphStyle("CustomH1", parent=styles["Heading1"],
                                  fontName=font_for_pdf, fontSize=16, spaceBefore=12, spaceAfter=6)
        h2_style = ParagraphStyle("CustomH2", parent=styles["Heading2"],
                                  fontName=font_for_pdf, fontSize=13, spaceBefore=10, spaceAfter=4)
        body_style = ParagraphStyle("CustomBody", parent=styles["Normal"],
                                    fontName=font_for_pdf, fontSize=10, leading=14, spaceAfter=6)

        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 6*mm))

        for block in parse_markdown_blocks(markdown):
            self._render_pdf_block(story, block, font_for_pdf, styles, title_style, h1_style, h2_style, body_style)

        try:
            doc.build(story)
            logger.info(f"PDF exported to {filepath}")
            return filepath
        except Exception as e:
            logger.error(f"PDF build failed: {e}")
            return None

    @staticmethod
    def _pdf_inline(text: str) -> str:
        """行内标记 → reportlab 支持的迷你 HTML 标签。"""
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
        text = re.sub(r"`(.+?)`", r"<font face='Courier'>\1</font>", text)
        return text

    def _render_pdf_block(self, story, block, font_for_pdf, styles,
                          title_style, h1_style, h2_style, body_style) -> None:
        """单个 Block → reportlab story 元素。"""
        if isinstance(block, Blank):
            story.append(Spacer(1, 3*mm))
        elif isinstance(block, Heading):
            style = {1: h1_style, 2: h2_style}.get(block.level, styles["Heading3"])
            story.append(Paragraph(self._pdf_inline(block.text), style))
        elif isinstance(block, MdTable):
            self._add_pdf_table(story, block, font_for_pdf)
        elif isinstance(block, HorizontalRule):
            story.append(Spacer(1, 4*mm))
        elif isinstance(block, MdImage):
            png = self._resolve_chart_image(block.ref)
            if png:
                try:
                    story.append(self._scaled_image(png))
                    story.append(Spacer(1, 4 * mm))
                except Exception as e:
                    logger.warning(f"pdf image failed ({png}): {e}")
                    story.append(Paragraph(f"（图表：{block.alt or '图'}，无法嵌入）", body_style))
            else:
                story.append(Paragraph(f"（图表：{block.alt or '图'}）", body_style))
        elif isinstance(block, ListItem):
            clean = self._pdf_inline(block.text)
            if block.ordered:
                # 历史行为：有序列表项按普通正文渲染，保留编号
                story.append(Paragraph(f"{block.marker} {clean}", body_style))
            else:
                story.append(Paragraph(f"• {clean}", body_style))
        else:  # MdParagraph
            clean = self._pdf_inline(block.text)
            if clean:
                story.append(Paragraph(clean, body_style))

    def _add_pdf_table(self, story, table: MdTable, font_name: str):
        """把解析后的表格渲染为 reportlab Table。"""
        if not table.header:
            return
        col_count = len(table.header)
        table_data = [table.header] + [r[:col_count] + [""] * (col_count - len(r)) for r in table.rows]
        style = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2f2f2")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fafafa")]),
        ])
        tbl = Table(table_data, style=style, hAlign="LEFT")
        story.append(tbl)
        story.append(Spacer(1, 4 * mm))

    def _export_html(self, markdown: str, title: str) -> str:
        """导出为 HTML 文件（用基本 Markdown→HTML 转换）。"""
        output_dir = _ensure_export_dir("html")
        filename = _make_filename(title, "html")
        filepath = os.path.join(output_dir, filename)

        html = self._md_to_html(markdown, title)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(html)
        logger.info(f"HTML exported to {filepath}")
        return filepath

    def _md_to_html(self, markdown: str, title: str) -> str:
        """将 Markdown 转为基本 HTML。"""
        html_lines = [
            "<!DOCTYPE html>",
            "<html><head><meta charset='utf-8'>",
            f"<title>{title}</title>",
            "<style>",
            "body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }",
            "h1 { border-bottom: 2px solid #333; padding-bottom: 8px; }",
            "h2 { border-bottom: 1px solid #ddd; padding-bottom: 4px; }",
            "table { border-collapse: collapse; width: 100%; margin: 10px 0; }",
            "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "th { background-color: #f2f2f2; }",
            "img { max-width: 100%; }",
            "</style></head><body>",
        ]
        state = {"in_table": False}

        def close_table():
            if state["in_table"]:
                html_lines.append("</tbody></table>")
                state["in_table"] = False

        for block in parse_markdown_blocks(markdown):
            if isinstance(block, Blank):
                close_table()
                html_lines.append("")
            elif isinstance(block, Heading):
                close_table()
                html_lines.append(f"<h{block.level}>{block.text}</h{block.level}>")
            elif isinstance(block, MdTable):
                if not state["in_table"]:
                    html_lines.append("<table><thead>")
                    state["in_table"] = True
                cleaned_header = [self._html_clean(c) for c in block.header]
                html_lines.append(f"<tr><th>{'</th><th>'.join(cleaned_header)}</th></tr>")
                html_lines.append("</thead><tbody>")
                for row in block.rows:
                    cleaned = [self._html_clean(c) for c in row]
                    html_lines.append(f"<tr><td>{'</td><td>'.join(cleaned)}</td></tr>")
            elif isinstance(block, HorizontalRule):
                close_table()
                html_lines.append("<hr>")
            elif isinstance(block, MdImage):
                close_table()
                png = self._resolve_chart_image(block.ref)
                src = ""
                if png:
                    try:
                        src = self._png_data_uri(png)
                    except Exception as e:
                        logger.warning(f"html image data-uri failed ({png}): {e}")
                        src = ""
                if src:
                    html_lines.append(f'<img src="{src}" alt="{self._html_clean(block.alt)}">')
                else:
                    html_lines.append(f"<p>（图表：{block.alt or '图'}）</p>")
            elif isinstance(block, ListItem):
                close_table()
                clean = self._html_clean(block.text)
                clean = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", clean)
                clean = re.sub(r"\*(.+?)\*", r"<em>\1</em>", clean)
                clean = re.sub(r"`(.+?)`", r"<code>\1</code>", clean)
                html_lines.append(f"<li>{clean}</li>")
            else:  # MdParagraph
                close_table()
                clean = self._html_clean(block.text)
                clean = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", clean)
                clean = re.sub(r"\*(.+?)\*", r"<em>\1</em>", clean)
                clean = re.sub(r"`(.+?)`", r"<code>\1</code>", clean)
                html_lines.append(f"<p>{clean}</p>")

        close_table()
        html_lines.append("</body></html>")
        return "\n".join(html_lines)

    @staticmethod
    def _html_clean(text: str) -> str:
        """HTML 转义。"""
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _make_filename(title: str, ext: str) -> str:
    """生成安全文件名。"""
    safe = re.sub(r"[^\w\s-]", "", title)
    safe = re.sub(r"[-\s]+", "_", safe.strip())[:50]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{safe}_{timestamp}.{ext}"
