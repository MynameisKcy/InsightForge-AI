import hashlib
import os.path
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader

from utils.logger_handler import logger



def get_file_md5_hex(filepath:str):
    if not os.path.exists(filepath):
        logger.error('File %s not exists!', filepath)
        return None
    if not os.path.isfile(filepath):
        logger.error('File %s is not a file!', filepath)
        return None
    md5_obj = hashlib.md5()
    chunk_size = 4096
    try:
        with open(filepath, 'rb') as f:
            while chunk := f.read(chunk_size):
                md5_obj.update(chunk)
        return md5_obj.hexdigest()
    except Exception as e:
        logger.error(e)
        return None

def listdir_with_allowed_type(path:str, allowed_types:tuple[str]):
    files = []
    if not os.path.isdir(path):
        logger.error('%s is not a directory!', path)
        return tuple()

    for f in os.listdir(path):
        if f.endswith(allowed_types):
            files.append(os.path.join(path, f))
    return tuple(files)

def pdf_loader(filepath:str,passwd:str=None) -> list[Document]:
    return PyPDFLoader(filepath,passwd).load()

def text_loader(filepath:str,passwd:str=None) -> list[Document]:
    return TextLoader(filepath, encoding=passwd or "utf-8").load()

def docx_loader(filepath:str,passwd:str=None) -> list[Document]:
    """加载 .docx 文档，按段落抽取文本，保留来源元数据。"""
    import docx
    doc = docx.Document(filepath)
    parts = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    # 表格内容也纳入
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(cells)
            if line:
                parts.append(line)
    content = "\n".join(parts)
    if not content.strip():
        return []
    return [Document(page_content=content, metadata={"source": filepath})]

def markdown_loader(filepath:str,passwd:str=None) -> list[Document]:
    """加载 markdown 文件（本质是文本，复用 TextLoader）。"""
    return TextLoader(filepath, encoding=passwd or "utf-8").load()
